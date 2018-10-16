# -*- coding: utf-8 -*-
'''
pip install python-docx
'''
import sys
reload(sys)
sys.setdefaultencoding("utf-8")
from docx import Document
from docx.shared import Inches, Pt, RGBColor

document = Document()

document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='IntenseQuote')

document.add_paragraph(
    'first item in unordered list', style='ListBullet'
)
document.add_paragraph(
    'first item in ordered list', style='ListNumber'
)

#document.add_picture('monty-truth.png', width=Inches(1.25))
document.add_paragraph()
paragraph = document.add_paragraph()
paragraph_format = paragraph.paragraph_format
paragraph_format.left_indent = Inches(0.5) ## left indent
paragraph_format.right_indent = Pt(24) ## right indent
paragraph_format.line_spacing = Pt(40)  ## line space


run = paragraph.add_run(u"run test 中文 paragraphparagraphparagraphp\
paragraph paragraph paragraph paragraph ")
font = run.font
font.name = u"华文仿宋"
font.size = Pt(20)
font.italic = True
font.bold = True
font.underline = True
font.color.rgb = RGBColor(0x42, 0x24, 0xE9)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
#for item in recordset:
#    row_cells = table.add_row().cells
#    row_cells[0].text = str(item.qty)
#    row_cells[1].text = str(item.id)
#    row_cells[2].text = item.desc

document.add_page_break()

document.save('demo.docx')
