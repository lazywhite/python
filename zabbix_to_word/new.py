'''
    get config by self-defined tags in docx

    tag format
        zabbix_graph|hostname|graphname

    only support docx, not doc
'''
from __future__ import print_function
import sys
reload(sys)
sys.setdefaultencoding('utf-8')
from docx import Document
from docx.shared import Inches, Pt, Cm
from lib.graph import Graph
from lib.zapi import ZAPI
from lib.config import get_config
import re

from pprint import pprint
from datetime import datetime

import pdb

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None



if __name__ == '__main__':
    now = datetime.now().strftime('%Y%m%d_%H%M')
    cfg = get_config("config.yml")
#    pprint(cfg)

    server = cfg.get("zabbix_server")
    username = cfg.get("zabbix_user")
    password = cfg.get("zabbix_pw")
    graph_path = cfg.get("graph_dir")
    period = cfg.get('period')
    stime = cfg.get('start_time')
    template_name = cfg.get('template_name')
    client_name = template_name.split('.')[0]
    template = cfg.get('template_dir') + template_name
    
    output_path = cfg.get('output_dir')
    output_name = client_name + '_' + now + '.docx'

    zp = ZAPI(server, username, password)
    gph = Graph(server, username, password)

#    doc = Document(unicode(template))
    doc = Document(template)
#    pdb.set_trace()
    for p in doc.paragraphs:
        print(p.text)
        if p.text.strip().startswith('zabbix_graph'):
            m=re.match('^zabbix_graph\|([^|]*)\|([^|]*)$', p.text.strip())
            val1, val2 = m.groups()
            hostname = val1.strip()
            graph = val2.strip()
            host_id = zp.get_host_id_by_name(hostname)
            graph_id = zp.get_graph_id(host_id, graph)
            graph_name = graph_path + '_'.join([hostname, graph_id, now, '.png'])
            gph.get_graph(graph_id,period, stime, graph_name)

            new = p.insert_paragraph_before()
            r = new.add_run()
            r.add_picture(graph_name, width=Cm(17))
            delete_paragraph(p)

                                
    doc.save(output_path + output_name)
