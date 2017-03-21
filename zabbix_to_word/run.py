from __future__ import print_function
import sys
reload(sys)
sys.setdefaultencoding('utf-8')
from docx import Document
from docx.shared import Inches, Pt
from lib.graph import Graph
from lib.zapi import ZAPI
from lib.config import get_config

from pprint import pprint
from datetime import datetime

#import pdb

if __name__ == '__main__':
    now = datetime.now().strftime('%Y%m%d_%H%M')
    cfg = get_config("config.yml")
#    pprint(cfg)

    server = cfg.get("zabbix_server")
    username = cfg.get("zabbix_user")
    password = cfg.get("zabbix_pw")
    graph_path = cfg.get("graph_dir")
    period = cfg.get('period')
    stime = cfg.get('start_time')
    template_name = cfg.get('template_name')
    client_name = template_name.split('.')[0]
    template = cfg.get('template_dir') + template_name
    
    output_path = cfg.get('output_dir')
    output_name = client_name + '_' + now + '.docx'
    doc = Document(template)
    doc.add_page_break()

    zp = ZAPI(server, username, password)
    gph = Graph(server, username, password)

    graphs = cfg.get("graphs")
    
    for section in graphs:
        for section_name, section_values in section.iteritems():
#            pdb.set_trace()
            doc.add_heading(section_name, level=1)
            for category in section_values:
                for category_name, category_value in category.iteritems():
                    doc.add_heading(category_name, level=2)
                    for host_graph in category_value:
                        for hostname, host_graphs in host_graph.iteritems():
                            doc.add_heading(hostname, level=3)
                            host_id = zp.get_host_id_by_name(hostname)
                            for graph in host_graphs:
                                graph_id = zp.get_graph_id(host_id, graph)
                                graph_name = graph_path + '_'.join([hostname, graph, now, '.png'])
                                doc.add_heading(graph, level=4)
                                gph.get_graph(graph_id,period, stime, graph_name)
                                doc.add_picture(graph_name, width=Inches(5.5))


                                
    doc.save(output_path + output_name)
